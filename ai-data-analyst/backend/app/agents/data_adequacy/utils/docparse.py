"""Document parsing utilities for various file formats."""

import os
import re
import logging
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import tempfile
import hashlib

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import pandas as pd
import textstat

from ..config import config

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of various document formats."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.csv', '.xlsx', '.xls'}
    
    def __init__(self):
        self.max_file_size = config.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate file before processing."""
        path = Path(file_path)
        
        # Check if file exists
        if not path.exists():
            return False, "File does not exist"
        
        # Check file size
        if path.stat().st_size > self.max_file_size:
            return False, f"File size exceeds {config.MAX_FILE_SIZE_MB}MB limit"
        
        # Check file extension
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file format: {path.suffix}"
        
        return True, "Valid file"
    
    def parse_file(self, file_path: str) -> Dict:
        """Parse a file and extract text and metadata."""
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            return {"success": False, "error": message}
        
        path = Path(file_path)
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                result = self._parse_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                result = self._parse_docx(file_path)
            elif extension == '.txt':
                result = self._parse_txt(file_path)
            elif extension in ['.csv', '.xlsx', '.xls']:
                result = self._parse_tabular(file_path)
            else:
                return {"success": False, "error": f"Parser not implemented for {extension}"}
            
            # Add common metadata
            result.update({
                "file_path": file_path,
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "file_extension": extension,
                "file_hash": self._compute_file_hash(file_path)
            })
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            return {"success": False, "error": f"Parsing failed: {str(e)}"}
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """Parse PDF file using PyMuPDF and pdfplumber as fallback."""
        text_content = []
        metadata = {}
        
        try:
            # Try PyMuPDF first (faster)
            doc = fitz.open(file_path)
            metadata.update({
                "page_count": doc.page_count,
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", "")
            })
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append({
                        "page": page_num + 1,
                        "text": text.strip(),
                        "type": "text"
                    })
                
                # Extract tables if present
                tables = page.find_tables()
                for table in tables:
                    table_data = table.extract()
                    if table_data:
                        text_content.append({
                            "page": page_num + 1,
                            "text": self._table_to_text(table_data),
                            "type": "table",
                            "raw_table": table_data
                        })
            
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying pdfplumber: {str(e)}")
            # Fallback to pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata["page_count"] = len(pdf.pages)
                    
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text:
                            text_content.append({
                                "page": i + 1,
                                "text": text.strip(),
                                "type": "text"
                            })
                        
                        # Extract tables
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                text_content.append({
                                    "page": i + 1,
                                    "text": self._table_to_text(table),
                                    "type": "table",
                                    "raw_table": table
                                })
                                
            except Exception as e2:
                return {"success": False, "error": f"Both PDF parsers failed: {str(e2)}"}
        
        full_text = "\n\n".join([content["text"] for content in text_content])
        
        return {
            "success": True,
            "text": full_text,
            "content_sections": text_content,
            "metadata": metadata,
            "language": self._detect_language(full_text),
            "readability_score": textstat.flesch_reading_ease(full_text) if full_text else 0
        }
    
    def _parse_docx(self, file_path: str) -> Dict:
        """Parse DOCX file."""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                table_text = self._docx_table_to_text(table)
                if table_text:
                    text_content.append(table_text)
            
            full_text = "\n\n".join(text_content)
            
            return {
                "success": True,
                "text": full_text,
                "content_sections": [{"text": full_text, "type": "document"}],
                "metadata": {
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables)
                },
                "language": self._detect_language(full_text),
                "readability_score": textstat.flesch_reading_ease(full_text) if full_text else 0
            }
            
        except Exception as e:
            return {"success": False, "error": f"DOCX parsing failed: {str(e)}"}
    
    def _parse_txt(self, file_path: str) -> Dict:
        """Parse plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                "success": True,
                "text": text,
                "content_sections": [{"text": text, "type": "text"}],
                "metadata": {
                    "character_count": len(text),
                    "word_count": len(text.split()),
                    "line_count": len(text.split('\n'))
                },
                "language": self._detect_language(text),
                "readability_score": textstat.flesch_reading_ease(text) if text else 0
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return {"success": False, "error": "Could not decode file with any encoding"}
            
            return {
                "success": True,
                "text": text,
                "content_sections": [{"text": text, "type": "text"}],
                "metadata": {"encoding_used": encoding},
                "language": self._detect_language(text),
                "readability_score": textstat.flesch_reading_ease(text) if text else 0
        
        except Exception as e:
            return {"success": False, "error": f"Text parsing failed: {str(e)}"}
    
    def _parse_tabular(self, file_path: str) -> Dict:
        """Parse tabular data (CSV, Excel)."""
        try:
            path = Path(file_path)
            
            if path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:  # Excel files
                df = pd.read_excel(file_path)
            
            # Convert DataFrame to text representation
            text_parts = []
            
            # Add column headers
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            
            # Add sample rows (first 10)
            text_parts.append("\nSample data:")
            text_parts.append(df.head(10).to_string(index=False))
            
            # Add data summary
            text_parts.append(f"\nData Summary:")
            text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            
            # Add column types
            text_parts.append("\nColumn Types:")
            for col, dtype in df.dtypes.items():
                text_parts.append(f"{col}: {dtype}")
            
            full_text = "\n".join(text_parts)
            
            return {
                "success": True,
                "text": full_text,
                "content_sections": [{"text": full_text, "type": "tabular"}],
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "data_types": df.dtypes.to_dict(),
                    "has_null_values": df.isnull().any().any(),
                    "null_counts": df.isnull().sum().to_dict()
                },
                "raw_data": df.to_dict('records')[:100],  # First 100 rows
                "language": "data",  # Special language for tabular data
                "readability_score": 100  # Tabular data is considered highly readable
            }
            
        except Exception as e:
            return {"success": False, "error": f"Tabular parsing failed: {str(e)}"}
    
    def _table_to_text(self, table_data: List[List]) -> str:
        """Convert table data to text representation."""
        if not table_data:
            return ""
        
        text_lines = []
        for row in table_data:
            if row and any(cell for cell in row if cell):  # Skip empty rows
                text_lines.append(" | ".join(str(cell) if cell else "" for cell in row))
        
        return "\n".join(text_lines)
    
    def _docx_table_to_text(self, table) -> str:
        """Convert DOCX table to text."""
        text_lines = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            if any(cells):  # Skip empty rows
                text_lines.append(" | ".join(cells))
        
        return "\n".join(text_lines)
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on common patterns."""
        if not text:
            return "unknown"
        
        # Simple heuristic - could be improved with proper language detection library
        text_sample = text[:1000].lower()
        
        english_indicators = ['the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one']
        english_count = sum(1 for word in english_indicators if word in text_sample)
        
        if english_count >= 3:
            return "english"
        
        # Add more language detection logic as needed
        return "unknown"
    
    def _compute_file_hash(self, file_path: str) -> str:
        """Compute SHA-256 hash of file."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def detect_ocr_artifacts(self, text: str) -> Tuple[bool, float]:
        """Detect if text contains OCR artifacts."""
        if not text:
            return False, 0.0
        
        # Patterns that indicate OCR issues
        ocr_patterns = [
            r'\b[a-zA-Z]{1,2}\b',  # Single or double character "words"
            r'[|!1Il]{2,}',  # Multiple pipes/exclamation marks (often confused by OCR)
            r'\d[a-zA-Z]\d',  # Number-letter-number patterns
            r'[^\w\s]{3,}',  # Long sequences of special characters
        ]
        
        total_matches = 0
        total_words = len(text.split())
        
        for pattern in ocr_patterns:
            matches = len(re.findall(pattern, text))
            total_matches += matches
        
        if total_words == 0:
            return False, 0.0
        
        artifact_ratio = total_matches / total_words
        has_artifacts = artifact_ratio > 0.05  # 5% threshold
        
        return has_artifacts, artifact_ratio
